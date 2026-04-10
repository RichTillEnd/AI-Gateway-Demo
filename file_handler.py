"""
檔案處理工具
處理上傳的圖片和文件
"""

import os
import base64
from PIL import Image
from io import BytesIO
from typing import Tuple, Optional
import mimetypes

# 建立上傳目錄
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 允許的檔案類型
ALLOWED_IMAGE_TYPES = {'.png', '.jpg', '.jpeg', '.webp', '.gif'}
ALLOWED_DOCUMENT_TYPES = {'.pdf', '.txt', '.md', '.docx', '.xlsx', '.csv'}
ALLOWED_CODE_TYPES = {'.py', '.js', '.java', '.cpp', '.html', '.css', '.json', '.xml'}

ALL_ALLOWED_TYPES = ALLOWED_IMAGE_TYPES | ALLOWED_DOCUMENT_TYPES | ALLOWED_CODE_TYPES

# 最大檔案大小 (20MB)
MAX_FILE_SIZE = 20 * 1024 * 1024


def get_file_category(filename: str) -> str:
    """
    根據副檔名判斷檔案類型
    
    返回: 'image', 'document', 'code', 或 'unknown'
    """
    ext = os.path.splitext(filename)[1].lower()
    
    if ext in ALLOWED_IMAGE_TYPES:
        return 'image'
    elif ext in ALLOWED_DOCUMENT_TYPES:
        return 'document'
    elif ext in ALLOWED_CODE_TYPES:
        return 'code'
    else:
        return 'unknown'


def validate_file(filename: str, file_size: int) -> Tuple[bool, str]:
    """
    驗證檔案是否符合規範
    
    返回: (是否通過, 錯誤訊息)
    """
    # 檢查檔案大小
    if file_size > MAX_FILE_SIZE:
        return False, f"檔案大小超過限制 (最大 {MAX_FILE_SIZE // 1024 // 1024}MB)"
    
    # 檢查副檔名
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALL_ALLOWED_TYPES:
        return False, f"不支援的檔案類型: {ext}"
    
    return True, ""


def save_upload_file(file_content: bytes, filename: str, user_id: int) -> str:
    """
    儲存上傳的檔案
    
    返回: 檔案路徑
    """
    # 建立用戶專屬目錄
    user_dir = os.path.join(UPLOAD_DIR, f"user_{user_id}")
    os.makedirs(user_dir, exist_ok=True)
    
    # 生成唯一檔名（加上時間戳避免重複）
    import time
    timestamp = int(time.time())
    name, ext = os.path.splitext(filename)
    safe_filename = f"{name}_{timestamp}{ext}"
    
    filepath = os.path.join(user_dir, safe_filename)
    
    # 儲存檔案
    with open(filepath, 'wb') as f:
        f.write(file_content)
    
    return filepath


def process_image_for_ai(filepath: str, max_size: int = 2000) -> str:
    """
    處理圖片供 AI 使用
    
    步驟:
    1. 壓縮大圖片
    2. 轉換成 base64
    
    返回: base64 編碼的圖片
    """
    try:
        # 開啟圖片
        img = Image.open(filepath)
        
        # 如果圖片太大，縮小它
        if max(img.size) > max_size:
            ratio = max_size / max(img.size)
            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        # 轉換成 RGB（去除 alpha 通道）
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background
        
        # 轉換成 base64
        buffered = BytesIO()
        img.save(buffered, format="JPEG", quality=85)
        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
        
        return img_base64
        
    except Exception as e:
        raise Exception(f"圖片處理失敗: {str(e)}")


def extract_text_from_file(filepath: str) -> Optional[str]:
    """
    從檔案中提取文字內容
    
    支援:
    - PDF
    - Word (.docx)
    - Excel (.xlsx)
    - 純文字 (.txt, .md, .py, .js 等)
    
    返回: 提取的文字或 None
    """
    ext = os.path.splitext(filepath)[1].lower()
    
    try:
        # 純文字檔案
        if ext in {'.txt', '.md', '.py', '.js', '.java', '.cpp', '.html', '.css', '.json', '.xml', '.csv'}:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        
        # PDF 檔案
        elif ext == '.pdf':
            return extract_text_from_pdf(filepath)
        
        # Word 檔案
        elif ext == '.docx':
            return extract_text_from_docx(filepath)
        
        # Excel 檔案
        elif ext == '.xlsx':
            return extract_text_from_xlsx(filepath)
        
        else:
            return None
            
    except Exception as e:
        print(f"提取文字失敗: {str(e)}")
        return None


def extract_text_from_pdf(filepath: str) -> str:
    """從 PDF 提取文字"""
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(filepath)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
        
    except Exception as e:
        raise Exception(f"PDF 讀取失敗: {str(e)}")


def extract_text_from_docx(filepath: str) -> str:
    """從 Word 文件提取文字"""
    try:
        from docx import Document
        
        doc = Document(filepath)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 也提取表格內容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        
        return "\n".join(text_parts)
        
    except Exception as e:
        raise Exception(f"Word 文件讀取失敗: {str(e)}")


def extract_text_from_xlsx(filepath: str) -> str:
    """從 Excel 提取文字"""
    try:
        from openpyxl import load_workbook
        
        wb = load_workbook(filepath, read_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"\n=== 工作表: {sheet_name} ===\n")
            
            for row in sheet.iter_rows(values_only=True):
                # 過濾掉空行
                row_values = [str(cell) if cell is not None else '' for cell in row]
                if any(row_values):
                    text_parts.append(' | '.join(row_values))
        
        return "\n".join(text_parts)
        
    except Exception as e:
        raise Exception(f"Excel 讀取失敗: {str(e)}")


def get_file_info(filepath: str) -> dict:
    """
    獲取檔案資訊
    
    返回: 包含檔案資訊的字典
    """
    filename = os.path.basename(filepath)
    file_size = os.path.getsize(filepath)
    mime_type, _ = mimetypes.guess_type(filepath)
    
    return {
        'filename': filename,
        'filepath': filepath,
        'file_size': file_size,
        'mime_type': mime_type or 'application/octet-stream',
        'file_type': get_file_category(filename)
    }


def cleanup_old_files(user_id: int, days: int = 30):
    """
    清理超過指定天數的舊檔案
    
    參數:
        user_id: 用戶 ID
        days: 保留天數
    """
    import time
    
    user_dir = os.path.join(UPLOAD_DIR, f"user_{user_id}")
    if not os.path.exists(user_dir):
        return
    
    current_time = time.time()
    cutoff_time = current_time - (days * 24 * 60 * 60)
    
    for filename in os.listdir(user_dir):
        filepath = os.path.join(user_dir, filename)
        if os.path.isfile(filepath):
            file_time = os.path.getmtime(filepath)
            if file_time < cutoff_time:
                try:
                    os.remove(filepath)
                    print(f"已刪除舊檔案: {filepath}")
                except Exception as e:
                    print(f"刪除檔案失敗: {str(e)}")


if __name__ == "__main__":
    # 測試功能
    print("📁 檔案處理工具測試")
    print(f"上傳目錄: {UPLOAD_DIR}")
    print(f"支援的圖片格式: {ALLOWED_IMAGE_TYPES}")
    print(f"支援的文件格式: {ALLOWED_DOCUMENT_TYPES}")
    print(f"支援的程式碼格式: {ALLOWED_CODE_TYPES}")
    print(f"最大檔案大小: {MAX_FILE_SIZE // 1024 // 1024}MB")
